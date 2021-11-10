import os
import PyPDF2
import pikepdf
import pdfplumber
import pandas as pd
import docx
import comtypes.client
import win32com.client

"""pdf破解相关函数"""


def pdf_Crack(sourcePath, savePath, paswrd=""):
    """
    去除没有打开口令的pdf加密。
    不能覆盖原文件，支持中文名称和路径。
    :param sourcePath: 要破解的pdf文件路径
    :param savePath: 生成的pdf文件路径
    :return:生成的pdf文件
    """
    # with pikepdf.open(sourcePath) as pdf:
    #     pdf.save(savePath)
    #     print('successful cracked')
    #     return savePath
    fp = open(sourcePath, "rb+")
    pdfFile = PyPDF2.PdfFileReader(fp)

    filepath, tempfilename = os.path.split(sourcePath)
    if pdfFile.isEncrypted:
        try:
            pdf = pikepdf.open(sourcePath, password=paswrd)
            pdf.save(savePath)
            return savePath
        except:
            return 0
    else:
        with pikepdf.open(sourcePath) as pdf:
            pdf.save(savePath)
            print('successful cracked')
            return savePath


def batch_pdf_Crack(folderPath):
    """
    将一个文件夹中的文件全部解密，在main.py目录下保存为原名字。
    支持带中文的文件名字。
    支持带中文的路径。
    :param filePath: 要解密的文件夹
    :param savePath: 文件保存的文件夹
    :return:
    """
    filelist = os.listdir(folderPath)
    # os.listdir() 方法用于返回指定的文件夹包含的文件或文件夹的名字的列表。
    for file in filelist:
        if file.endswith(".pdf") and ("~$" not in file):
            filePath = folderPath + "\\" + file
            with pikepdf.open(filePath) as pdf:
                pdf.save(file)
    print('successful cracked')


"""pdf内容提取相关函数"""


def getTable(sourcePath, bp=0, ep=-1):
    """
    从pdf文件中取得表格。
    :param sourcePath:
    :param bp:beginPage
    :param ep:endPage
    :return :所取得表格的dataframe形式
    """
    pdf = pdfplumber.open(sourcePath)
    if ep == -1:
        ep = len(pdf.pages) - 1
    page = pdf.pages[bp - 1]
    result = page.extract_table()
    result = pd.DataFrame(result[:])
    for i in range(bp, ep):
        page = pdf.pages[i]
        table = page.extract_table()
        table = pd.DataFrame(table[:])
        result = pd.concat([result, table], axis=0)

    # print(type(table))
    # print(table)
    return result


def writeExcel(tabledf, savePath, sheet='Sheet1'):
    """
    将表格dataframe保存到excel文件中。
    :param tabledf: dataframe变量
    :param savePath: 要写入的excel文件
    :param sheet: 要写入的sheet
    :return:
    """
    # 使用to_excel写入xls文件要安装openpyxl库
    tabledf.to_excel(savePath, sheet)


def getText(sourcePath, bp=0, ep=-1):
    """
    从pdf文件中取得文本。
    :param sourcePath:
    :param bp:beginPage
    :param ep:endPage
    :return :文本字符串
    """
    pdf = pdfplumber.open(sourcePath)
    if ep == -1:
        ep = len(pdf.pages) - 1
    s = ""
    for i in range(bp - 1, ep):
        page = pdf.pages[i]
        table = page.extract_text()
        s += '\r\n' + table
    return s


def writeText(s, savePath):
    """
    将获取的文本写入docx或txt文件。
    :param s:
    :param savePath:
    :return:
    """
    if savePath.split('.')[-1] == 'txt':
        with open(savePath, 'a') as f:
            f.write('\n\r' + s)
    else:
        try:
            doc = docx.Document(savePath)
        except:
            doc = docx.Document()
        doc.add_paragraph(s)
        doc.save(savePath)


# ppt转pdf相关文档地址：https://pythonhosted.org/comtypes/
def init_powerpoint():
    powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
    # print(help(powerpoint))
    return powerpoint


def ppt_to_pdf(powerpoint, inputFileName, outputFileName, formatType=32):
    if outputFileName[-3:] != 'pdf':
        outputFileName = outputFileName + ".pdf"
    deck = powerpoint.Presentations.Open(inputFileName)
    # comtypes.POINTER（Presentations）实例的Open（…）方法
    deck.SaveAs(outputFileName, formatType)  # formatType = 32 for ppt to pdf
    # 这个formatType的取值见
    # https://docs.microsoft.com/en-us/office/vba/api/powerpoint.ppsaveasfiletype
    deck.Close()


def ppt2PDF_in_folder(powerpoint, folder):
    files = os.listdir(folder)
    pptfiles = [f for f in files if f.endswith((".ppt", ".pptx"))]
    for pptfile in pptfiles:
        fullpath = os.path.join(folder, pptfile)
        ppt_to_pdf(powerpoint, fullpath, fullpath)


def batch_ppt2PDF(powerpoint, fileList, savePath=None):
    if not savePath:
        saveList = [i.partition('.')[0] for i in fileList]
    else:
        saveList = [savePath + '\\' + i.split('\\')[-1] for i in fileList]
        saveList = [i.partition('.')[0] for i in saveList]
    for i in range(len(fileList)):
        ppt_to_pdf(powerpoint, fileList[i], saveList[i])


def word2pdf(doc_name, pdf_name):
    """
    :word文件转pdf
    :param doc_name word文件名称
    :param pdf_name 转换后pdf文件名称
    """
    if pdf_name[-3:] != 'pdf':
        pdf_name = pdf_name + ".pdf"
    try:
        word = win32com.client.DispatchEx("Word.Application")
        if os.path.exists(pdf_name):
            os.remove(pdf_name)
        worddoc = word.Documents.Open(doc_name, ReadOnly=1)
        worddoc.SaveAs(pdf_name, FileFormat=17)
        worddoc.Close()
        return pdf_name
    except:
        return 1


def batch_word2PDF(fileList, savePath=None):
    if not savePath:
        saveList = [i.partition('.')[0] for i in fileList]
    else:
        saveList = [savePath + '\\' + i.split('\\')[-1] for i in fileList]
        saveList = [i.partition('.')[0] for i in saveList]
    for i in range(len(fileList)):
        word2pdf(fileList[i], saveList[i])
